import os
import winreg
import ctypes

CARPETAS_IGNORADAS = {
    "windows", "program files", "program files (x86)", "programdata",
    "$recycle.bin", "system volume information", "appdata",
    "windowsapps", "msocache", "recovery", "perflogs",
    "node_modules", "__pycache__", "venv", "$sysreset"
}

EXTENSIONES_RELEVANTES = {
    ".docx", ".doc", ".pdf", ".txt", ".xlsx", ".xls", ".pptx", ".ppt", ".odt", ".csv",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp",
    ".mp4", ".mkv", ".avi", ".mov", ".mp3", ".wav", ".flac",
    ".exe", ".lnk",
    ".zip", ".rar", ".7z",
}

# Extensiones que SÍ se pueden abrir y leer como texto plano sin parser especial
EXTENSIONES_TEXTO_PLANO = {
    ".txt", ".csv", ".md", ".log", ".json", ".xml",
    ".py", ".js", ".html", ".css", ".ini", ".cfg", ".yaml", ".yml"
}

FILE_ATTRIBUTE_HIDDEN = 0x2
FILE_ATTRIBUTE_SYSTEM = 0x4
DRIVE_REMOVABLE = 2
DRIVE_FIXED = 3

class FileManager:
    @staticmethod
    def _leer_carpeta_registro(nombre_valor, fallback):
        """Lee la ruta real de una carpeta conocida de Windows desde el registro."""
        try:
            clave = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders"
            )
            valor, _ = winreg.QueryValueEx(clave, nombre_valor)
            winreg.CloseKey(clave)
            return os.path.expandvars(valor)
        except Exception:
            return fallback

    @staticmethod
    def obtener_carpetas_usuario():
        """Devuelve las rutas reales de las carpetas principales."""
        home = os.path.expanduser("~")
        return [
            FileManager._leer_carpeta_registro("Desktop", os.path.join(home, "Desktop")),
            FileManager._leer_carpeta_registro("Personal", os.path.join(home, "Documents")),
            FileManager._leer_carpeta_registro("{374DE290-123F-4565-9164-39C4925E467B}", os.path.join(home, "Downloads")),
            FileManager._leer_carpeta_registro("My Pictures", os.path.join(home, "Pictures")),
            FileManager._leer_carpeta_registro("My Video", os.path.join(home, "Videos")),
        ]

    @staticmethod
    def obtener_unidades_disponibles():
        """Obtiene letras de discos fijos o extraíbles."""
        unidades = []
        try:
            bitmask = ctypes.windll.kernel32.GetLogicalDrives()
            for i in range(26):
                if not (bitmask & (1 << i)):
                    continue
                letra = f"{chr(65 + i)}:\\"
                tipo = ctypes.windll.kernel32.GetDriveTypeW(letra)
                if tipo in (DRIVE_REMOVABLE, DRIVE_FIXED):
                    unidades.append(letra)
        except Exception:
            pass
        return unidades

    @staticmethod
    def _es_carpeta_ignorada(nombre_carpeta):
        return nombre_carpeta.lower() in CARPETAS_IGNORADAS or nombre_carpeta.startswith(".")

    @staticmethod
    def _es_archivo_oculto_o_sistema(ruta_completa):
        try:
            attrs = ctypes.windll.kernel32.GetFileAttributesW(ruta_completa)
            if attrs == -1:
                return False
            return bool(attrs & (FILE_ATTRIBUTE_HIDDEN | FILE_ATTRIBUTE_SYSTEM))
        except Exception:
            return False

    @staticmethod
    def buscar_archivos(carpetas_raiz, limite=2000):
        """Busca archivos relevantes en las carpetas dadas."""
        resultados = []
        vistos = set()
        for raiz in carpetas_raiz:
            if not os.path.exists(raiz):
                continue
            for carpeta_actual, subcarpetas, archivos in os.walk(raiz):
                subcarpetas[:] = [d for d in subcarpetas if not FileManager._es_carpeta_ignorada(d)]

                for nombre_archivo in archivos:
                    ext = os.path.splitext(nombre_archivo)[1].lower()
                    if ext not in EXTENSIONES_RELEVANTES:
                        continue

                    ruta_completa = os.path.join(carpeta_actual, nombre_archivo)
                    ruta_normalizada = os.path.normcase(os.path.abspath(ruta_completa))
                    
                    if ruta_normalizada in vistos:
                        continue
                    vistos.add(ruta_normalizada)

                    if FileManager._es_archivo_oculto_o_sistema(ruta_completa):
                        continue

                    resultados.append(ruta_completa)
                    if len(resultados) >= limite:
                        return resultados
        return resultados
    @staticmethod
    def abrir_archivo(ruta):
        """Abre un archivo con su programa predeterminado en Windows."""
        try:
            if os.path.exists(ruta):
                os.startfile(ruta)
                return True, f"Abriendo {os.path.basename(ruta)}"
            return False, "El archivo no existe."
        except Exception as e:
            return False, f"No pude abrir el archivo: {e}"

    @staticmethod
    def leer_archivo(ruta):
        """Lee el contenido de un archivo, usando el parser adecuado según su extensión."""
        if not os.path.exists(ruta):
            return False, "El archivo no existe."

        ext = os.path.splitext(ruta)[1].lower()

        try:
            if ext == ".docx":
                return FileManager._leer_docx(ruta)
            elif ext == ".pdf":
                return FileManager._leer_pdf(ruta)
            elif ext in (".xlsx", ".xls"):
                return FileManager._leer_xlsx(ruta)
            elif ext in EXTENSIONES_TEXTO_PLANO:
                with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
                    return True, f.read()
            else:
                return False, f"Todavía no sé leer el contenido de archivos {ext}."
        except ModuleNotFoundError as e:
            libreria = str(e).split("'")[1] if "'" in str(e) else str(e)
            return False, f"Falta instalar una librería para leer {ext}: pip install {libreria}"
        except Exception as e:
            return False, f"No pude leer el archivo: {e}"

    @staticmethod
    def _leer_docx(ruta):
        """Extrae texto (párrafos + tablas) de un Word .docx."""
        from docx import Document

        doc = Document(ruta)
        partes = [p.text for p in doc.paragraphs if p.text.strip()]

        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text.strip():
                        partes.append(celda.text)

        texto = "\n".join(partes)
        if not texto.strip():
            return False, "El documento Word está vacío o no tiene texto extraíble."
        return True, texto

    @staticmethod
    def _leer_pdf(ruta):
        """Extrae texto de un PDF, página por página."""
        from pypdf import PdfReader

        lector = PdfReader(ruta)
        partes = []
        for pagina in lector.pages:
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                partes.append(texto_pagina)

        texto = "\n".join(partes)
        if not texto.strip():
            return False, "No pude extraer texto de este PDF (puede ser un escaneo/imagen sin OCR)."
        return True, texto

    @staticmethod
    def _leer_xlsx(ruta):
        """Extrae el contenido de todas las hojas de un Excel."""
        from openpyxl import load_workbook

        wb = load_workbook(ruta, data_only=True, read_only=True)
        partes = []
        for nombre_hoja in wb.sheetnames:
            hoja = wb[nombre_hoja]
            partes.append(f"--- Hoja: {nombre_hoja} ---")
            for fila in hoja.iter_rows(values_only=True):
                valores = [str(v) for v in fila if v is not None]
                if valores:
                    partes.append(" | ".join(valores))

        texto = "\n".join(partes)
        if not texto.strip():
            return False, "El archivo Excel está vacío."
        return True, texto

    @staticmethod
    def escribir_archivo(ruta, contenido):
        """Sobrescribe un archivo (usar con cuidado)."""
        try:
            with open(ruta, 'w', encoding='utf-8') as f:
                f.write(contenido)
            return True, "Archivo guardado exitosamente."
        except Exception as e:
            return False, f"No pude escribir en el archivo: {e}"